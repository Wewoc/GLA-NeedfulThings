#!/usr/bin/env python3
"""
Step 6 (optional) — chats_to_word.py
Converts sorted chat JSON files to Word documents (.docx).

Requirements:
  pip install python-docx

Usage:
  python "Step 6 - optional - chats_to_word.py"
"""

import json
from datetime import datetime
from config import SORTED_DIR, WORD_EXPORT_DIR

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("✗ python-docx nicht installiert — pip install python-docx")
    raise SystemExit(1)

def format_ts(ts):
    if not ts: return "Kein Zeitstempel"
    try:
        return datetime.fromisoformat(str(ts).replace("Z", "+00:00")).strftime("%d.%m.%Y %H:%M")
    except:
        return str(ts)

def convert(path):
    try:
        data     = json.loads(path.read_text(encoding="utf-8"))
        title    = data.get("title") or data.get("name") or path.stem
        messages = data.get("messages") or data.get("chat_messages") or data.get("conversation")
        if not messages or not isinstance(messages, list):
            return False

        doc = Document()
        doc.add_heading(str(title), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Datei: {path.name}").italic = True
        doc.add_paragraph("-" * 30)

        for msg in messages:
            role    = str(msg.get("sender") or msg.get("role") or "").lower()
            content = str(msg.get("text") or msg.get("content") or "")
            ts      = format_ts(msg.get("created_at") or msg.get("timestamp"))
            p       = doc.add_paragraph()
            if any(x in role for x in ["human", "user", "mensch"]):
                run = p.add_run(f"👤 DU ({ts})")
                run.font.color.rgb = RGBColor(0, 51, 102)
            else:
                run = p.add_run(f"🤖 ASSISTANT ({ts})")
                run.font.color.rgb = RGBColor(0, 102, 51)
            run.bold = True
            doc.add_paragraph(content)
            doc.add_paragraph("." * 10).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(WORD_EXPORT_DIR / f"{path.stem}.docx")
        return True
    except Exception as e:
        print(f"  ✗ {path.name}: {e}")
        return False

def main():
    WORD_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    json_files = sorted([f for f in SORTED_DIR.glob("*.json") if f.name != "merged_chats.json"])
    print(f"Garmin Local Archive — JSON → Word")
    print(f"  {len(json_files)} Dateien gefunden")
    done = sum(1 for f in json_files if convert(f) and print(f"  ✓ {f.name}") is None)
    print(f"\n  ✓ {done} Word-Dateien erstellt in: {WORD_EXPORT_DIR.resolve()}")

if __name__ == "__main__":
    main()
